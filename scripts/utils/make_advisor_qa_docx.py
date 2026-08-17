# -*- coding: utf-8 -*-
"""교수님 첨삭(0815.hwp) 질문 5건에 대한 답변 문서 생성."""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_COLOR_INDEX, WD_ALIGN_PARAGRAPH

PROF  = RGBColor(0xC0,0x00,0x00)   # 빨강 — 교수님 원문 인용
AINEW = RGBColor(0x00,0x00,0xC0)   # 파랑 — AI 신규작성
GREY  = RGBColor(0x60,0x60,0x60)

d = Document()
st = d.styles['Normal']
st.font.name = '맑은 고딕'; st.font.size = Pt(10)

def H(t, sz=14, color=None):
    p = d.add_paragraph(); r = p.add_run(t)
    r.bold = True; r.font.size = Pt(sz)
    if color: r.font.color.rgb = color
    return p

def P(t, color=None, bold=False, sz=10, ul=False, hl=False, indent=None):
    p = d.add_paragraph()
    if indent: p.paragraph_format.left_indent = Cm(indent)
    r = p.add_run(t); r.font.size = Pt(sz); r.bold = bold
    if color: r.font.color.rgb = color
    if ul: r.font.underline = True
    if hl: r.font.highlight_color = WD_COLOR_INDEX.YELLOW
    return p

def QA(qno, sec, quote, note=None):
    H(f'{qno}. [{sec}] 교수님 표시', 12)
    p = d.add_paragraph(); p.paragraph_format.left_indent = Cm(0.5)
    r = p.add_run(f'"{quote}"'); r.font.color.rgb = PROF; r.font.size = Pt(10)
    if note:
        pp = d.add_paragraph(); pp.paragraph_format.left_indent = Cm(0.5)
        rr = pp.add_run(note); rr.font.size = Pt(9); rr.font.color.rgb = GREY; rr.italic = True

def TBL(rows, widths=None):
    t = d.add_table(rows=1, cols=len(rows[0])); t.style = 'Light Grid Accent 1'
    for i,h in enumerate(rows[0]):
        c = t.rows[0].cells[i]; c.text = ''
        r = c.paragraphs[0].add_run(h); r.bold = True; r.font.size = Pt(9)
    for row in rows[1:]:
        cells = t.add_row().cells
        for i,v in enumerate(row):
            cells[i].text = ''
            r = cells[i].paragraphs[0].add_run(str(v)); r.font.size = Pt(9)
    d.add_paragraph()
    return t

# ── 표지 ──
ti = d.add_paragraph(); ti.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = ti.add_run('EdgeGround-VLA — 교수님 첨삭 질문 답변'); r.bold = True; r.font.size = Pt(16)
sub = d.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run('대상: 초안 0815.hwp   |   작성 2026-08-17   |   질문 5건 전부 실측 완료')
r.font.size = Pt(9); r.font.color.rgb = GREY
d.add_paragraph()

P('※ 색상 규칙 — 빨강: 교수님 표시 원문 / 파랑: 신규 작성 답변 / 노란 밑줄: 재확인 필요',
  color=GREY, sz=8)
d.add_paragraph()

H('요약', 13)
TBL([
 ['#','위치','질문','답변 요지'],
 ['Q1','§5.3','"방향 대조 학습으로 학습하고 고정한다" (표시)','용어 정의 + 5.4절 예고 삽입. 용어도 재검토 권고'],
 ['Q2','§5.3','"1.128M" (표시)','0.262M + 0.866M 산식과 전체 대비 0.25% 병기'],
 ['Q3','§5.4 그림','image_proj — train/val 정확도','train 96.10% / val 94.09% (격차 2.01%p) — 신규 측정'],
 ['Q4','§5.5 그림','Action Head — train/val 정확도','train 91.44% / val 74.04% (격차 17.40%p)'],
 ['Q5','§5.5 그림','Action Head — train/val loss','train 0.169 / val 1.953 (epoch 225)'],
])

d.add_page_break()

# ── Q1 + Q2 ──
H('Q1 · Q2 — §5.3 두 표시는 같은 문단이므로 한 번에 수정', 13)
QA('Q1','§5.3','방향 대조 학습으로 학습하고 고정한다.')
QA('Q2','§5.3','1.128M',
   '※ HWP 파일 파싱 결과 밑줄 속성은 검출되지 않았으나(문자 서식상 0건), 표시된 것으로 간주하고 답변함.')
d.add_paragraph()

H('수정 전 (현재 원문)', 11)
P('학습 분할은 에피소드 단위로 15%를 validation으로 떼어냈으며(SPLIT_SEED=42), 프레임 단위가 아니라 '
  '에피소드 단위로 분할하였다. 학습 대상은 Stage 1에서 image_proj(0.262M)를 방향 대조 학습으로 학습하고 '
  '고정한다. 그리고 Stage 2에서 그 위에 MLP 헤드(0.866M)만 학습하는 구조이다. 전체 학습 파라미터는 '
  '1.128M이다. OWL-v2와 Kosmos-2 비전 타워는 두 단계 모두 frozen 상태로 유지된다.', indent=0.5)
d.add_paragraph()

H('수정 후 (제안)', 11)
p = d.add_paragraph(); p.paragraph_format.left_indent = Cm(0.5)
r = p.add_run('학습 분할은 에피소드 단위로 15%를 validation으로 떼어냈으며(SPLIT_SEED=42), 프레임 단위가 아니라 '
  '에피소드 단위로 분할하였다. 학습은 두 단계로 나뉜다. Stage 1에서는 image_proj(0.262M)를 방향 대조 학습, '
  '즉 이미지에서 추출한 특징을 방향을 서술하는 5개 문장의 텍스트 특징과 대조하여 정렬시키는 학습(5.4절에서 '
  '상술)으로 사전 학습한 뒤 그 가중치를 고정한다. Stage 2에서는 고정된 image_proj 위에서 MLP 행동 '
  '헤드(0.866M)만 학습한다. 따라서 역전파로 갱신되는 파라미터는 두 단계를 합쳐 1.128M(= 0.262M + 0.866M)이며, '
  '이는 추론에 참여하는 전체 파라미터 0.46B의 약 0.25%에 해당한다. OWL-v2(0.155B)와 Kosmos-2 비전 '
  '타워(0.303B)는 두 단계 모두 frozen 상태로 유지되어 학습되지 않는다.')
r.font.color.rgb = AINEW; r.font.size = Pt(10)
d.add_paragraph()

H('무엇이 바뀌었는가', 11)
for t in [
 '① "방향 대조 학습"에 그 자리에서 한 줄 정의를 붙이고 "(5.4절에서 상술)"로 예고했습니다. '
 '원문에도 바로 다음 문단에 이미 쉬운 설명이 있으나, 표시된 지점에서 그 설명이 곧 나온다는 신호가 없었습니다.',
 '② 1.128M의 산식을 괄호로 명시했습니다(= 0.262M + 0.866M).',
 '③ "전체 파라미터의 약 0.25%"를 추가했습니다. 1.128M이 무엇에 대비해 작은지 기준이 없으면 '
 '경량성이 전달되지 않으며, 경량화가 본 논문의 주제이므로 이 비율이 핵심입니다.',
 '④ frozen 백본의 파라미터 수(0.155B·0.303B)를 병기해 위 비율의 분모를 검증 가능하게 했습니다.',
]:
    P(t, color=AINEW, indent=0.5)
d.add_paragraph()

P('각주 권장 — 엄밀히는 Stage 1에서 text_proj(2048→256, 0.525M)도 함께 갱신되지만, 배포 시 추론 서버가 '
  '체크포인트에서 image_proj 키만 읽으므로 배포 모델의 학습 파라미터로는 1.128M이 맞습니다.',
  color=AINEW, indent=0.5, ul=True, hl=True)
d.add_paragraph()

H('참고 — "방향 대조 학습" 용어 자체의 재검토 권고', 11)
P('"방향 대조 학습"은 표준 용어가 아니며, 엄밀히는 배치 내 negative 샘플링이 없는 고정 5-클래스 분류이므로 '
  '대조 학습(contrastive learning)의 정의에도 정확히 부합하지 않습니다. 문헌 표준 표현은 아래와 같습니다.',
  color=AINEW, indent=0.5)
TBL([
 ['표현','근거 (원문 확인)'],
 ['text-derived classifier',
  'CLIP(Radford et al., 2021) §3.1.2: "the text encoder is a hypernetwork which generates the '
  'weights of a linear classifier based on the text specifying the visual concepts that the classes represent"'],
 ['cosine similarity + temperature scaling',
  'CLIP §3.1.2: "The cosine similarity of these embeddings is then calculated, scaled by a temperature '
  'parameter τ, and normalized into a probability distribution via a softmax."  → 본 연구 구현과 1:1 대응'],
 ['class prototype',
  'CLIP 본문에는 없는 용어(전수 검색 0건). 후속 연구 용어 — CLIP-S⁴(arXiv:2305.01040) '
  'Figure 2: "Class Prototypes derived from CLIP"'],
])
P('제안 표현: "텍스트 앵커를 클래스 프로토타입으로 사용하는 방향 분류 학습(text-derived class prototypes)". '
  '단 CLIP은 추론 시 분류기가 고정인 반면 본 연구는 text_proj도 학습하므로 "fixed"라는 수식은 붙이지 '
  '않는 것이 정확합니다.', color=AINEW, indent=0.5)

d.add_page_break()

# ── Q3 ──
H('Q3 — image_proj 학습 곡선: Training과 validation의 정확도', 13)
QA('Q3','§5.4 그림','Training과 validation의 정확도는?')
d.add_paragraph()
TBL([
 ['구분','정확도','프레임 수','에피소드'],
 ['Training','96.10%','13,114','180'],
 ['Validation','94.09%','3,485','45'],
 ['격차','+2.01%p','—','—'],
])
P('검증 정확도 94.09%는 체크포인트에 저장된 값과 일치합니다. train-val 격차가 2.01%p에 불과하여 '
  'image_proj는 과적합이 거의 없습니다.', color=AINEW, indent=0.5)
P('※ 학습 스크립트가 매 epoch validation만 평가하고 train 정확도는 기록하지 않아, 저장된 체크포인트로 '
  '동일 분할(seed=42, 에피소드 단위)을 복원해 별도 측정하였습니다(2026-08-17).',
  color=GREY, sz=9, indent=0.5)
d.add_paragraph()

# ── Q4 Q5 ──
H('Q4 · Q5 — Action Head: Training/validation 정확도 및 loss', 13)
QA('Q4','§5.5 그림','Training과 validation의 정확도는?')
QA('Q5','§5.5 그림','Training과 validation의 loss 값은?')
d.add_paragraph()
TBL([
 ['시점','train acc','val acc','train loss','val loss'],
 ['배포 체크포인트 (epoch 225)','91.44%','74.04%','0.169','1.953'],
 ['val loss 최저 (epoch 20)','72.29%','67.13%','0.668','0.926'],
 ['최종 (epoch 299)','92.24%','73.51%','0.146','2.040'],
])
import os
FIG='docs/v5/figures/action_head_mlp_training_curve.png'
if os.path.exists(FIG):
    d.add_picture(FIG, width=Cm(16))
    P('Action Head(MLP) 학습 곡선 — Accuracy(좌) / Loss(우), seed=0, 300 epoch',
      color=GREY, sz=8, indent=0.5)
d.add_paragraph()

H('함께 서술해야 할 사항 — 과적합', 11)
P('train 91.44% 대 val 74.04%로 17.40%p의 격차가 있으며, validation loss는 epoch 20 이후 '
  '0.926에서 1.953까지 상승합니다. 즉 행동 헤드에서는 과적합이 관찰됩니다. 다음 세 가지를 함께 '
  '서술하는 것이 적절합니다.', color=AINEW, indent=0.5)
for t in [
 '① 체크포인트 선택 기준은 validation accuracy 단독이며, 분류 과제에서 실제로 쓰이는 지표가 '
 'argmax 정확도이기 때문입니다. validation loss가 상승하는데 정확도는 유지되는 현상은 모델이 '
 '틀린 예측에 과신하게 되지만 argmax 순위는 뒤집히지 않는 알려진 패턴입니다.',
 '② 데이터가 225 에피소드(16,599 프레임)로 작아 과적합이 구조적으로 불가피합니다.',
 '③ 그럼에도 실기 100회 시험에서 95%의 목표 도달 성공률을 확인하였으며, 이는 validation 지표만으로 '
 '실제 주행 성능을 판단할 수 없음을 보여줍니다.',
]:
    P(t, color=AINEW, indent=0.8)
d.add_paragraph()

H('"왜 epoch 225인가" — 코드 확인 결과', 11)
for t in [
 '① 학습 스크립트에는 validation loss를 계산하는 코드가 없습니다. loss를 보고 선택한 것이 아니라 '
 '처음부터 validation accuracy만이 기준이었습니다.',
 '② validation은 25 epoch마다만 평가하므로 후보는 0·25·50·…·275·299의 13개뿐이며, 그중 225가 '
 '최고였습니다. 따라서 위 표의 epoch 20은 실제 학습에서 후보가 아니었습니다(5 epoch 간격 재현 '
 '측정에서만 관측되는 지점).',
 '③ 비교가 "이상(>=)"이므로 동점 시 나중 epoch이 채택되는 편향이 있습니다.',
]:
    P(t, color=AINEW, indent=0.8)
d.add_paragraph()
P('한계로 함께 밝힐 사항 — loss 기준 조기종료 모델이 실기에서 더 우수한지는 검증하지 않았으며, '
  'validation 집합을 체크포인트 선택과 성능 보고에 함께 사용하고 있습니다(별도 test 집합 없음).',
  color=AINEW, indent=0.5, ul=True, hl=True)

d.add_page_break()

# ── 본문 추가 제안 ──
H('추가 제안 — §5.5 말미 삽입 문단 (Stage 1 vs Stage 2 일반화 격차)', 13)
P('Q3과 Q4의 답변을 종합하면 두 단계의 일반화 양상이 뚜렷하게 다르므로, 이를 본문에 서술하면 '
  '과적합 지적에 대한 방어 논거가 됩니다.', color=GREY, sz=9, indent=0.5)
p = d.add_paragraph(); p.paragraph_format.left_indent = Cm(0.5)
r = p.add_run('Stage 1과 Stage 2의 일반화 양상은 뚜렷하게 다르다. image_proj(Stage 1)는 학습 정확도 '
 '96.10%(13,114 프레임) 대 검증 정확도 94.09%(3,485 프레임)로 그 격차가 2.01%p에 불과하나, 행동 '
 '헤드(Stage 2)는 학습 정확도 91.44% 대 검증 정확도 74.04%로 격차가 17.40%p에 이른다. 즉 과적합은 '
 '파이프라인 전반의 문제가 아니라 Stage 2에 국한된 현상이다. 이는 목표물의 방향을 5단계로 판별하는 '
 '과제(Stage 1)는 현재의 데이터 규모에서 충분히 일반화되는 반면, 그 표현 위에서 8개 행동 클래스를 '
 '판별하는 과제(Stage 2)는 동일한 데이터 규모에 비해 난이도가 높다는 것을 시사한다. 행동 헤드의 검증 '
 '손실이 학습 후반에 상승하는 것 역시 같은 원인으로 해석된다. 다만 이러한 격차가 실제 주행 성능을 '
 '제한하지는 않았으며, 실로봇 100회 시험에서 95%의 목표 도달 성공률을 확인하였다(IV장).')
r.font.color.rgb = AINEW; r.font.size = Pt(10)
d.add_paragraph()

H('참고 — 첨삭본에서 발견한 그림 번호 충돌', 11)
P('§5.4~5.5에서 서로 다른 세 그림이 모두 "Figure 8"로 표기되어 있습니다 '
  '(image_proj 학습 곡선 / Action Head - Accuracy / Action Head - Loss). 번호 재부여가 필요합니다.',
  color=AINEW, indent=0.5, ul=True, hl=True)
TBL([
 ['현재','권장','내용'],
 ['Figure 8','Figure 5','image_proj 학습 곡선 (§5.4)'],
 ['Figure 8','Figure 6','Action Head — Accuracy (§5.5)'],
 ['Figure 8','Figure 7','Action Head — Loss (§5.5)'],
 ['Fig 9','Figure 8','val 혼동행렬 (뒤로 밀림)'],
])
P('※ Action Head 곡선을 Accuracy/Loss 두 그림으로 분리하는 경우의 안입니다. 하나의 2-패널 그림으로 '
  '합치는 경우에는 Figure 6 하나로 두고 이후 번호를 한 칸씩 당기면 됩니다.', color=GREY, sz=9, indent=0.5)

out='docs/01.paper/EdgeGround-VLA_교수님질문_답변_0817.docx'
d.save(out)
print('saved:', out)
